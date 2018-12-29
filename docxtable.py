#!/usr/bin/env python3
# coding=utf8

from __future__ import print_function
import requests
from docx.api import Document
import os
import tempfile
from sys import exit
from datetime import date, timedelta, datetime
import re

DOMAIN = 'http://m-ontv.net'
MOSAD = 288
CLASS = 'יב 4'

# find table
ontv = requests.get('{}/indexEmbed.asp?theMosad={}&theIndex=index1'.format(DOMAIN, MOSAD)).text
page = requests.get('{}/loadMessagesAjax.asp?themosad={}&screen=2&side=both&displayMode=static'.format(DOMAIN, MOSAD)).text
fileID = re.search(r'data-fileID="(\d+)"', page).group(1)
docx_filename = requests.get('{}/getFileMessageDetailsAjax.asp?fileID={}'.format(DOMAIN, fileID)).text
docx_url = DOMAIN + '/' + docx_filename[:docx_filename.find(',')].replace('\\', '/')
# feed to docxlib
tf, fn = tempfile.mkstemp()
os.write(tf, requests.get(docx_url).content)
doc = Document(fn)
title = doc.paragraphs[0].text.strip()
day_to_check = date.today() + timedelta(days=1-(datetime.now().hour < 15))
if title[title.rfind(' ')+1:] != day_to_check.strftime('%-d.%-m.%y'):
    print(title[title.rfind(' ')+1:], tomorrow.strftime('%-d.%-m.%y'), date.today().strftime('%-d.%-m.%y'), datetime.now().hour < 15)
    print('טבלה לא מעודכנת')
    exit(1)
table = doc.tables[0]

# find class
row = None
for r in table.rows:
    if r.cells[0].text == CLASS:
        row = r
        break
else:
    print('Illegal class chosen.')
    exit(2)

print(title)
# parse changes
changes = [c.text for c in row.cells]
hr = 1
while hr < len(changes):
    cell = changes[hr]
    start_hr = hr
    while hr < len(changes) - 1 and changes[hr + 1] == changes[start_hr]:
        hr += 1

    if cell == '':
        outstr = 'כרגיל'
    elif all(x == '/' for x in cell):
        outstr = 'פוצץ'
    else:
        outstr = cell
    if hr == start_hr:
        print('שעה {} {}'.format(hr, outstr))
    else:
        print('שעות {}-{} {}'.format(start_hr, hr, outstr))
    hr += 1
